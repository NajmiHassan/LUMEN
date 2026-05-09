"""Export research markdown to Word (.docx) and LaTeX for Overleaf and print."""

from __future__ import annotations

import io
import re
import urllib.parse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _split_inline_bold(text: str) -> list[tuple[str, bool]]:
    """Split text into (fragment, is_bold) segments for **bold** spans."""
    parts: list[tuple[str, bool]] = []
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            parts.append((text[pos : m.start()], False))
        parts.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False))
    return parts if parts else [(text, False)]


def latex_escape(text: str) -> str:
    """Escape LaTeX special characters in plain text."""
    out: list[str] = []
    for ch in text:
        if ch == "\\":
            out.append(r"\textbackslash{}")
        elif ch == "&":
            out.append(r"\&")
        elif ch == "%":
            out.append(r"\%")
        elif ch == "$":
            out.append(r"\$")
        elif ch == "#":
            out.append(r"\#")
        elif ch == "_":
            out.append(r"\_")
        elif ch == "{":
            out.append(r"\{")
        elif ch == "}":
            out.append(r"\}")
        elif ch == "~":
            out.append(r"\textasciitilde{}")
        elif ch == "^":
            out.append(r"\textasciicircum{}")
        else:
            out.append(ch)
    return "".join(out)


def _latex_inline(text: str) -> str:
    """Apply bold handling and escape remainder of fragment."""
    chunks: list[str] = []
    for fragment, bold in _split_inline_bold(text):
        esc = latex_escape(fragment)
        chunks.append(r"\textbf{" + esc + "}" if bold else esc)
    return "".join(chunks)


def parse_blocks(markdown: str) -> list[tuple[str, str]]:
    """
    Parse markdown into simple blocks: title, h1–h3, para, bullet, numbered.

    Headings follow # / ## / ### / #### at line start.
    """
    blocks: list[tuple[str, str]] = []
    lines = markdown.splitlines()
    buf: list[str] = []

    def flush_para() -> None:
        nonlocal buf
        if buf:
            blocks.append(("para", " ".join(buf)))
            buf = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_para()
            continue

        if stripped.startswith("#### "):
            flush_para()
            blocks.append(("h3", stripped[5:].strip()))
            continue
        if stripped.startswith("### "):
            flush_para()
            blocks.append(("h2", stripped[4:].strip()))
            continue
        if stripped.startswith("## "):
            flush_para()
            blocks.append(("h1", stripped[3:].strip()))
            continue
        if stripped.startswith("# "):
            flush_para()
            blocks.append(("title", stripped[2:].strip()))
            continue

        bullet_match = re.match(r"^[-*•]\s+", stripped)
        if bullet_match:
            flush_para()
            blocks.append(("bullet", stripped[bullet_match.end() :].strip()))
            continue

        num_match = re.match(r"^\d+\.\s+", stripped)
        if num_match:
            flush_para()
            blocks.append(("numbered", stripped[num_match.end() :].strip()))
            continue

        buf.append(stripped)

    flush_para()
    return blocks


def _add_formatted_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for fragment, bold in _split_inline_bold(text):
        run = p.add_run(fragment)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        if bold:
            run.bold = True


def markdown_to_docx_bytes(markdown: str) -> bytes:
    """
    Build a research-style Word document with title and heading hierarchy.
    """
    doc = Document()
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    blocks = parse_blocks(markdown)

    for kind, text in blocks:
        if not text:
            continue
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "Calibri"
            continue

        if kind == "h1":
            h = doc.add_heading(text, level=1)
            for para_run in h.runs:
                para_run.font.name = "Calibri"
                para_run.font.size = Pt(14)
            continue

        if kind == "h2":
            h = doc.add_heading(text, level=2)
            for para_run in h.runs:
                para_run.font.name = "Calibri"
                para_run.font.size = Pt(12)
            continue

        if kind == "h3":
            h = doc.add_heading(text, level=3)
            for para_run in h.runs:
                para_run.font.name = "Calibri"
            continue

        if kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            for fragment, bold in _split_inline_bold(text):
                run = p.add_run(fragment)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                if bold:
                    run.bold = True
            continue

        if kind == "numbered":
            p = doc.add_paragraph(style="List Number")
            for fragment, bold in _split_inline_bold(text):
                run = p.add_run(fragment)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                if bold:
                    run.bold = True
            continue

        # para
        _add_formatted_paragraph(doc, text)

    for kind, text in blocks:
        if kind == "title" and text:
            doc.core_properties.title = text
            break

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def markdown_to_latex(markdown: str) -> str:
    """
    Convert parsed markdown blocks into a standalone LaTeX article suitable for Overleaf.
    """
    blocks = parse_blocks(markdown)
    title = "Research Report"
    out: list[str] = []

    i = 0
    while i < len(blocks):
        kind, text = blocks[i]
        if not text:
            i += 1
            continue

        if kind == "title":
            title = text
            i += 1
            continue

        if kind == "h1":
            out.append(f"\\section{{{_latex_inline(text)}}}")
            i += 1
            continue

        if kind == "h2":
            out.append(f"\\subsection{{{_latex_inline(text)}}}")
            i += 1
            continue

        if kind == "h3":
            out.append(f"\\subsubsection{{{_latex_inline(text)}}}")
            i += 1
            continue

        if kind == "bullet":
            items: list[str] = []
            while i < len(blocks) and blocks[i][0] == "bullet" and blocks[i][1]:
                items.append(f"\\item {_latex_inline(blocks[i][1])}")
                i += 1
            out.append("\\begin{itemize}")
            out.extend(items)
            out.append("\\end{itemize}")
            out.append("")
            continue

        if kind == "numbered":
            items = []
            while i < len(blocks) and blocks[i][0] == "numbered" and blocks[i][1]:
                items.append(f"\\item {_latex_inline(blocks[i][1])}")
                i += 1
            out.append("\\begin{enumerate}")
            out.extend(items)
            out.append("\\end{enumerate}")
            out.append("")
            continue

        if kind == "para":
            out.append(_latex_inline(text))
            out.append("")
            i += 1
            continue

        i += 1

    body = "\n".join(out)
    body = re.sub(r"\n{3,}", "\n\n", body).strip()

    doc = f"""\\documentclass[11pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[T1]{{fontenc}}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{hyperref}}
\\usepackage{{parskip}}
\\hypersetup{{colorlinks=true, linkcolor=blue, urlcolor=blue, citecolor=blue}}

\\title{{{latex_escape(title)}}}
\\author{{AI Research Assistant}}
\\date{{\\today}}

\\begin{{document}}
\\maketitle
\\tableofcontents
\\newpage

{body}

\\end{{document}}
"""
    return doc


def overleaf_new_project_url() -> str:
    """Open Overleaf new blank project; user uploads/pastes the generated .tex file."""
    return "https://www.overleaf.com/project/new"


def overleaf_edit_url_from_latex(latex_source: str) -> str:
    """
    Try opening Overleaf with the document preloaded via the public `/docs` snip flow.

    If the URL would be too long for browsers or Overleaf limits, fall back to
    opening a blank project (user uploads the downloaded ``main.tex``).
    """
    url, _ = overleaf_open_info(latex_source)
    return url


def overleaf_open_info(latex_source: str) -> tuple[str, bool]:
    """
    Returns (url, snippet_supported).

    When ``snippet_supported`` is False, the URL opens a blank Overleaf project;
    the user should upload the downloaded `.tex` file instead.
    """
    base = "https://www.overleaf.com/docs"
    encoded = urllib.parse.quote(latex_source, safe="")
    candidate = f"{base}?snip={encoded}"
    if len(candidate) > 4000:
        return overleaf_new_project_url(), False
    return candidate, True
